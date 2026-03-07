"""
document_builder.py
Assembles the final training document (.docx) in step-by-step manual style.
Includes: cover page, TOC, sections with screenshots, steps, tips, warnings, quiz.
"""

import io
import logging
import requests
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

# ── Brand colours ─────────────────────────────────────────────────────────────
BLUE       = RGBColor(0x00, 0x71, 0xC5)   # Intel blue
DARK_BLUE  = RGBColor(0x00, 0x3E, 0x7E)
LIGHT_GREY = RGBColor(0xF5, 0xF5, 0xF5)
DARK_GREY  = RGBColor(0x44, 0x44, 0x44)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
ORANGE     = RGBColor(0xE6, 0x5C, 0x00)   # warning colour
GREEN      = RGBColor(0x10, 0x7C, 0x10)   # tip colour


def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _add_horizontal_rule(doc: Document, color: str = "0071C5"):
    """Add a coloured horizontal rule paragraph."""
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    return p


def _add_page_break(doc: Document):
    doc.add_page_break()


def _fetch_image_bytes(url: str, blob_client=None) -> bytes | None:
    """
    Fetch image bytes. If url is a blob path (no http), fetch from Azure blob storage.
    Otherwise fetch from URL directly.
    """
    try:
        # Blob path saved during captions stage (e.g. "PM/frames/frame_001.jpg")
        if blob_client and not url.startswith("http"):
            blob = blob_client.get_blob_client(container="intermediate", blob=url)
            return blob.download_blob().readall()
        # Regular URL
        resp = requests.get(url, timeout=20)
        resp.raise_for_status()
        return resp.content
    except Exception as e:
        logger.warning(f"Could not fetch image from {url}: {e}")
        return None


def _setup_styles(doc: Document):
    """Configure custom paragraph styles."""
    styles = doc.styles

    # Section heading
    try:
        h1 = styles["Heading 1"]
    except KeyError:
        h1 = styles.add_style("Heading 1", WD_STYLE_TYPE.PARAGRAPH)
    h1.font.name  = "Calibri"
    h1.font.size  = Pt(16)
    h1.font.bold  = True
    h1.font.color.rgb = DARK_BLUE

    # Step heading
    try:
        h2 = styles["Heading 2"]
    except KeyError:
        h2 = styles.add_style("Heading 2", WD_STYLE_TYPE.PARAGRAPH)
    h2.font.name  = "Calibri"
    h2.font.size  = Pt(12)
    h2.font.bold  = True
    h2.font.color.rgb = BLUE

    # Normal body
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


# ── COVER PAGE ────────────────────────────────────────────────────────────────

def _build_cover_page(doc: Document, title: str, overview: str):
    # Blue header band
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    _set_cell_bg(cell, "0071C5")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after  = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TRAINING MANUAL")
    run.font.name  = "Calibri"
    run.font.size  = Pt(12)
    run.font.bold  = True
    run.font.color.rgb = WHITE

    doc.add_paragraph()

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(title)
    run.font.name  = "Calibri"
    run.font.size  = Pt(28)
    run.font.bold  = True
    run.font.color.rgb = DARK_BLUE

    doc.add_paragraph()
    _add_horizontal_rule(doc)

    # Overview
    p_overview = doc.add_paragraph()
    p_overview.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_overview.add_run(overview)
    run.font.name  = "Calibri"
    run.font.size  = Pt(12)
    run.font.color.rgb = DARK_GREY

    doc.add_paragraph()
    doc.add_paragraph()

    # Date
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_date.add_run(f"Generated: {datetime.utcnow().strftime('%B %d, %Y')}")
    run.font.name  = "Calibri"
    run.font.size  = Pt(10)
    run.font.color.rgb = DARK_GREY

    _add_page_break(doc)


# ── PREREQUISITES ─────────────────────────────────────────────────────────────

def _build_prerequisites(doc: Document, prerequisites: list[str]):
    if not prerequisites:
        return
    h = doc.add_heading("Prerequisites", level=1)
    h.runs[0].font.color.rgb = DARK_BLUE
    for prereq in prerequisites:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(prereq).font.size = Pt(11)
    doc.add_paragraph()


# ── SECTION ───────────────────────────────────────────────────────────────────

def _build_section(
    doc: Document,
    section: dict,
    section_number: int,
    captioned_frames: list[dict],
    blob_client=None,
):
    content = section.get("content", {})

    # ── Section header ────────────────────────────────────
    h = doc.add_heading(f"Section {section_number}: {section['title']}", level=1)
    h.runs[0].font.color.rgb = DARK_BLUE
    _add_horizontal_rule(doc)

    # ── Objective box ─────────────────────────────────────
    obj_table = doc.add_table(rows=1, cols=1)
    obj_table.style = "Table Grid"
    obj_cell = obj_table.cell(0, 0)
    _set_cell_bg(obj_cell, "E6F2FB")
    p = obj_cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    label = p.add_run("🎯  Learning Objective:  ")
    label.font.bold  = True
    label.font.color.rgb = DARK_BLUE
    label.font.size  = Pt(11)
    obj_run = p.add_run(section.get("objective", ""))
    obj_run.font.size = Pt(11)
    doc.add_paragraph()

    # ── Screenshot ────────────────────────────────────────
    screenshot_ts  = section.get("screenshot_timestamp", -1)
    best_frame     = _find_closest_frame(captioned_frames, screenshot_ts)
    if best_frame:
        img_bytes = _fetch_image_bytes(best_frame["url"], blob_client=blob_client)
        if img_bytes:
            try:
                img_stream = io.BytesIO(img_bytes)
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p_img.add_run()
                run.add_picture(img_stream, width=Inches(5.5))

                # Caption
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap.add_run(
                    f"Figure {section_number}: {best_frame.get('caption', '')}"
                )
                cap_run.font.size  = Pt(9)
                cap_run.font.italic = True
                cap_run.font.color.rgb = DARK_GREY
                doc.add_paragraph()
            except Exception as e:
                logger.warning(f"Could not insert screenshot: {e}")

    # ── Introduction ──────────────────────────────────────
    intro = content.get("introduction", "")
    if intro:
        p = doc.add_paragraph()
        p.add_run(intro).font.size = Pt(11)
        doc.add_paragraph()

    # ── Steps ─────────────────────────────────────────────
    steps = content.get("steps", [])
    if steps:
        h2 = doc.add_heading("Steps", level=2)
        h2.runs[0].font.color.rgb = BLUE

    for step in steps:
        step_num  = step.get("step_number", "")
        instr     = step.get("instruction", "")
        detail    = step.get("detail")
        tip       = step.get("tip")
        warning   = step.get("warning")

        # Step instruction
        p_step = doc.add_paragraph()
        p_step.paragraph_format.left_indent  = Inches(0.2)
        p_step.paragraph_format.space_before = Pt(6)
        num_run = p_step.add_run(f"Step {step_num}:  ")
        num_run.font.bold  = True
        num_run.font.color.rgb = BLUE
        num_run.font.size  = Pt(11)
        instr_run = p_step.add_run(instr)
        instr_run.font.size = Pt(11)

        # Detail
        if detail:
            p_det = doc.add_paragraph()
            p_det.paragraph_format.left_indent = Inches(0.5)
            r = p_det.add_run(detail)
            r.font.size  = Pt(10)
            r.font.color.rgb = DARK_GREY

        # Tip box
        if tip:
            tip_tbl = doc.add_table(rows=1, cols=1)
            tip_tbl.style = "Table Grid"
            tip_cell = tip_tbl.cell(0, 0)
            _set_cell_bg(tip_cell, "E8F5E9")
            tp = tip_cell.paragraphs[0]
            tp.paragraph_format.left_indent = Inches(0.1)
            tr = tp.add_run(f"💡  Tip:  ")
            tr.font.bold  = True
            tr.font.color.rgb = GREEN
            tr.font.size  = Pt(10)
            tr2 = tp.add_run(tip)
            tr2.font.size = Pt(10)
            doc.add_paragraph()

        # Warning box
        if warning:
            warn_tbl = doc.add_table(rows=1, cols=1)
            warn_tbl.style = "Table Grid"
            warn_cell = warn_tbl.cell(0, 0)
            _set_cell_bg(warn_cell, "FFF3E0")
            wp = warn_cell.paragraphs[0]
            wp.paragraph_format.left_indent = Inches(0.1)
            wr = wp.add_run(f"⚠️  Warning:  ")
            wr.font.bold  = True
            wr.font.color.rgb = ORANGE
            wr.font.size  = Pt(10)
            wr2 = wp.add_run(warning)
            wr2.font.size = Pt(10)
            doc.add_paragraph()

    # ── Summary ───────────────────────────────────────────
    summary = content.get("summary", "")
    if summary:
        doc.add_paragraph()
        p_sum = doc.add_paragraph()
        r = p_sum.add_run("Summary:  ")
        r.font.bold  = True
        r.font.size  = Pt(11)
        r2 = p_sum.add_run(summary)
        r2.font.size = Pt(11)

    # ── Key takeaways ─────────────────────────────────────
    takeaways = content.get("key_takeaways", [])
    if takeaways:
        doc.add_paragraph()
        h_kt = doc.add_heading("Key Takeaways", level=2)
        h_kt.runs[0].font.color.rgb = BLUE
        for kt in takeaways:
            p_kt = doc.add_paragraph(style="List Bullet")
            p_kt.add_run(kt).font.size = Pt(11)

    _add_page_break(doc)


# ── QUIZ ──────────────────────────────────────────────────────────────────────

def _build_quiz(doc: Document, questions: list[dict]):
    if not questions:
        return

    h = doc.add_heading("Knowledge Check", level=1)
    h.runs[0].font.color.rgb = DARK_BLUE
    _add_horizontal_rule(doc)

    p_intro = doc.add_paragraph()
    r = p_intro.add_run(
        "Test your understanding of the material covered in this training document."
    )
    r.font.size = Pt(11)
    doc.add_paragraph()

    for i, q in enumerate(questions, 1):
        # Question
        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(8)
        qr = p_q.add_run(f"Q{i}.  {q.get('question', '')}")
        qr.font.bold  = True
        qr.font.size  = Pt(11)

        # Options
        for opt in q.get("options", []):
            p_opt = doc.add_paragraph(style="List Bullet")
            p_opt.paragraph_format.left_indent = Inches(0.3)
            p_opt.add_run(opt).font.size = Pt(11)

        # Answer (greyed out for trainer copy)
        ans = q.get("correct_answer", "")
        if ans:
            p_ans = doc.add_paragraph()
            p_ans.paragraph_format.left_indent = Inches(0.3)
            ar = p_ans.add_run(f"Answer: {ans}")
            ar.font.size  = Pt(10)
            ar.font.color.rgb = DARK_GREY
            ar.font.italic = True

        doc.add_paragraph()


# ── MAIN BUILD FUNCTION ───────────────────────────────────────────────────────

def build_document(
    structure: dict,
    captioned_frames: list[dict],
    quiz_questions: list[dict],
    output_path: str | None = None,
    blob_client=None,
) -> io.BytesIO:
    """
    Build the complete training document.
    Returns a BytesIO object containing the .docx file.
    Also saves to output_path if provided.
    """
    doc = Document()
    _setup_styles(doc)

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    title        = structure.get("title", "Training Document")
    overview     = structure.get("overview", "")
    prerequisites = structure.get("prerequisites", [])
    sections     = structure.get("sections", [])

    logger.info(f"Building document: '{title}' with {len(sections)} sections.")

    # Cover page
    _build_cover_page(doc, title, overview)

    # Prerequisites
    _build_prerequisites(doc, prerequisites)

    # Sections
    for i, section in enumerate(sections, 1):
        logger.info(f"  Building section {i}: {section['title']}")
        _build_section(doc, section, i, captioned_frames, blob_client=blob_client)

    # Quiz
    _build_quiz(doc, quiz_questions)

    # Save
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    if output_path:
        with open(output_path, "wb") as f:
            f.write(buf.read())
        buf.seek(0)
        logger.info(f"Document saved to {output_path}")

    logger.info("Document build complete.")
    return buf


# ── HELPERS ───────────────────────────────────────────────────────────────────

def _find_closest_frame(frames: list[dict], target_ts: float) -> dict | None:
    """Find the frame whose timestamp is closest to target_ts."""
    if not frames:
        return None
    if target_ts < 0:
        return frames[0]
    return min(frames, key=lambda f: abs(f.get("timestamp", 0) - target_ts))
